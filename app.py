import streamlit as st
import re
import time
import io
import os
from typing import Dict, Any, Optional, Tuple, List
from dataclasses import dataclass, field
from datetime import datetime
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import markdown
try:
    from xhtml2pdf import pisa
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False
import hashlib

# ==============================================================================
# SECTION 1: CONFIGURATION & CONSTANTS (DENGAN FALLBACK MODEL ASLI)
# ==============================================================================

@dataclass
class DocumentConfig:
    """Centralized configuration for document generation"""
    MARGIN_INCHES: float = 1.0
    # FALLBACK MODELS SESUAI VERSI ASLI ANDA
    FALLBACK_MODELS: List[str] = field(default_factory=lambda: [
        "gemini-2.5-flash",
        "gemini-2.5-pro",
        "gemini-2.5-flash-lite"
    ])
    MAX_OUTPUT_TOKENS: int = 8192
    TEMPERATURE: float = 0.1
    TOP_P: float = 0.95
    TOP_K: int = 40
    MAX_CHAT_HISTORY: int = 50
    MAX_FILE_SIZE_MB: int = 10
    SUPPORTED_EXTENSIONS: Tuple[str, ...] = ('.vtt', '.txt')
    
    def __post_init__(self):
        # MEMPERTAHANKAN FALLBACK MODELS ASLI ANDA
        if self.FALLBACK_MODELS is None:
            self.FALLBACK_MODELS = [
                "models/gemini-3.5-flash",       
                "models/gemini-3.1-flash-lite",  
                "models/gemini-2.5-flash",       
                "models/gemini-2.5-pro",         
                "models/gemini-2.5-flash-lite"   
            ]

@dataclass
class UIMessages:
    """Centralized UI messages"""
    PROCESSING = "🤖 Memproses data..."
    ANALYZING = "🔍 Menganalisis transkrip..."
    GENERATING = "📝 Menghasilkan notulen..."
    BUILDING_DOC = "📄 Membangun dokumen..."
    SUCCESS = "✅ Berhasil!"
    ERROR_API = "❌ API Key tidak valid atau quota habis"
    ERROR_FILE = "❌ Gagal memproses file"
    ERROR_PDF = "❌ Gagal generate PDF"

# Initialize config
config = DocumentConfig()

# ==============================================================================
# SECTION 2: UTILITY FUNCTIONS (IMPROVED)
# ==============================================================================

def validate_file(uploaded_file) -> Tuple[bool, str]:
    """Validate uploaded file size and type"""
    if uploaded_file is None:
        return False, "Tidak ada file"
    
    # Check file size
    if uploaded_file.size > config.MAX_FILE_SIZE_MB * 1024 * 1024:
        return False, f"File terlalu besar. Maksimal {config.MAX_FILE_SIZE_MB}MB"
    
    # Check extension
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    if file_ext not in config.SUPPORTED_EXTENSIONS:
        return False, f"Format tidak didukung. Gunakan: {', '.join(config.SUPPORTED_EXTENSIONS)}"
    
    return True, "OK"

def safe_file_read(uploaded_file) -> str:
    """Safely read file with encoding fallback"""
    try:
        return uploaded_file.getvalue().decode("utf-8")
    except UnicodeDecodeError:
        try:
            return uploaded_file.getvalue().decode("latin-1")
        except:
            return uploaded_file.getvalue().decode("utf-8", errors="ignore")

def process_vtt_text(vtt_text: str) -> str:
    """Improved VTT processing with better error handling"""
    if not vtt_text:
        return ""
    
    try:
        # Handle bytes input
        if isinstance(vtt_text, bytes):
            vtt_text = vtt_text.decode('utf-8', errors='ignore')
        
        # Clean VTT content - mempertahankan logic asli tapi lebih robust
        cleaned = re.sub(r"\d{2}:\d{2}:\d{2}\.\d{3} --> .*\n?", "", vtt_text)
        cleaned = re.sub(r"WEBVTT.*\n?", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"X-.*\n?", "", cleaned)  # Hapus header VTT kustom
        cleaned = re.sub(r"NOTE.*\n?", "", cleaned)  # Hapus notes
        
        # Remove empty lines and strip
        lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
        
        # Remove duplicate consecutive lines
        unique_lines = []
        for line in lines:
            if not unique_lines or line != unique_lines[-1]:
                unique_lines.append(line)
        
        return '\n'.join(unique_lines)
    except Exception as e:
        st.error(f"Error processing VTT: {str(e)}")
        return vtt_text

def clean_cell_text(text: str) -> str:
    """Clean and format cell text for Word documents"""
    if not text:
        return ""
    
    # Handle various newline formats
    text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    text = text.replace('&nbsp;', ' ')
    
    # Remove HTML tags but preserve structure
    text = re.sub(r'<[^>]+>', '', text)
    
    # Clean extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

def get_content_hash(content: str) -> str:
    """Generate hash for content caching"""
    return hashlib.md5(content.encode()).hexdigest()

# ==============================================================================
# SECTION 3: DOCUMENT GENERATION (IMPROVED)
# ==============================================================================

def create_word_document(content: str) -> Optional[io.BytesIO]:
    """Improved Word document generation with better table handling"""
    try:
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(config.MARGIN_INCHES)
        
        # Add title
        title = doc.add_heading('Notulen Rapat', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        timestamp_para.add_run(f"Dibuat: {datetime.now().strftime('%d/%m/%Y %H:%M')}").italic = True
        
        doc.add_paragraph()
        
        lines = content.split('\n')
        in_table = False
        table_data = []
        
        for line in lines:
            line_strip = line.strip()
            
            # Detect table rows (MARKDOWN TABLE)
            if line_strip.startswith('|') and line_strip.endswith('|'):
                if not in_table:
                    in_table = True
                    table_data = []
                
                # Skip separator rows (|---|---|)
                if re.match(r'^\|[-\s|]+\|$', line_strip):
                    continue
                
                cells = [cell.strip() for cell in line_strip.strip('|').split('|')]
                table_data.append(cells)
            
            else:
                # Process non-table content
                if in_table and table_data:
                    # Create table
                    if table_data:
                        rows = len(table_data)
                        cols = max(len(row) for row in table_data)
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        for i, row in enumerate(table_data):
                            for j in range(min(cols, len(row))):
                                clean_text = clean_cell_text(row[j])
                                if clean_text:
                                    # Handle multi-line text in cells
                                    if '\n' in clean_text:
                                        lines_in_cell = clean_text.split('\n')
                                        for k, cell_line in enumerate(lines_in_cell):
                                            if k == 0:
                                                table.cell(i, j).paragraphs[0].add_run(cell_line)
                                            else:
                                                table.cell(i, j).add_paragraph().add_run(cell_line)
                                    else:
                                        table.cell(i, j).text = clean_text
                        
                        doc.add_paragraph()  # Add space after table
                    in_table = False
                    table_data = []
                
                # Process headings and text
                if line_strip.startswith('# ') and line_strip != "# Notulen Rapat":
                    doc.add_heading(line_strip[2:], level=1)
                elif line_strip.startswith('## '):
                    doc.add_heading(line_strip[3:], level=2)
                elif line_strip.startswith('### '):
                    doc.add_heading(line_strip[4:], level=3)
                elif line_strip.startswith('- ') or line_strip.startswith('* '):
                    doc.add_paragraph(line_strip[2:], style='List Bullet')
                elif line_strip.startswith('**') and line_strip.endswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(line_strip[2:-2]).bold = True
                elif line_strip:
                    # Bold text handling
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*)', line_strip)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            p.add_run(part[2:-2]).bold = True
                        else:
                            p.add_run(part)
        
        # Handle trailing table
        if in_table and table_data:
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            for i, row in enumerate(table_data):
                for j in range(min(cols, len(row))):
                    table.cell(i, j).text = clean_cell_text(row[j])
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        return None

@st.cache_data(ttl=3600)  # Cache for 1 hour
def create_pdf_document_cached(content: str) -> Optional[bytes]:
    """Cached PDF generation"""
    return create_pdf_document(content)

def create_pdf_document(content: str) -> Optional[bytes]:
    """Improved PDF conversion with better error handling"""
    try:
        # Convert markdown to HTML with proper extensions
        html_content = markdown.markdown(
            content, 
            extensions=['tables', 'nl2br', 'fenced_code']
        )
        
        # Professional styling for PDF
        full_html = f"""
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                @page {{ 
                    size: A4; 
                    margin: 2cm;
                    @bottom-center {{
                        content: "Halaman " counter(page) " dari " counter(pages);
                        font-size: 9pt;
                        font-family: Helvetica, Arial, sans-serif;
                    }}
                }}
                body {{ 
                    font-family: Helvetica, Arial, sans-serif; 
                    font-size: 11pt; 
                    line-height: 1.5; 
                    color: #333;
                }}
                h1 {{ 
                    text-align: center; 
                    font-size: 18pt; 
                    margin-bottom: 20px;
                    color: #1a1a1a;
                }}
                h2 {{ 
                    font-size: 14pt; 
                    margin-top: 20px; 
                    margin-bottom: 10px;
                    color: #2c3e50;
                    border-bottom: 1px solid #ddd;
                    padding-bottom: 5px;
                }}
                h3 {{ 
                    font-size: 12pt; 
                    margin-top: 15px; 
                    margin-bottom: 8px;
                    color: #34495e;
                }}
                table {{ 
                    width: 100%; 
                    border-collapse: collapse; 
                    margin-top: 10px; 
                    margin-bottom: 15px;
                    page-break-inside: avoid;
                }}
                th, td {{ 
                    border: 1px solid #999; 
                    padding: 8px; 
                    vertical-align: top;
                    text-align: left;
                }}
                th {{ 
                    background-color: #f5f5f5; 
                    font-weight: bold;
                }}
                ul, ol {{ 
                    margin-top: 5px;
                    margin-bottom: 10px;
                    padding-left: 20px;
                }}
                li {{
                    margin-bottom: 3px;
                }}
                .footer {{
                    text-align: center;
                    font-size: 8pt;
                    color: #666;
                    margin-top: 20px;
                }}
                .mnev-watermark {{
                    position: fixed;
                    opacity: 0.05;
                    font-size: 60pt;
                    text-align: center;
                    width: 100%;
                    top: 40%;
                    pointer-events: none;
                }}
            </style>
        </head>
        <body>
            <div class="mnev-watermark">MNEV Intelligence</div>
            {html_content}
            <div class="footer">
                Dokumen ini digenerate oleh MNEV Intelligence pada {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
            </div>
        </body>
        </html>
        """
        
        result_file = io.BytesIO()
        pisa_status = pisa.CreatePDF(
            io.StringIO(full_html), 
            dest=result_file,
            encoding='utf-8'
        )
        
        if pisa_status.err:
            st.error(f"PDF generation error: {pisa_status.err}")
            return None
        
        result_file.seek(0)
        return result_file.getvalue()
    
    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return None

# ==============================================================================
# SECTION 4: AI CORE ENGINE (DENGAN FALLBACK MODEL ASLI ANDA)
# ==============================================================================

def generate_with_fallback(
    prompt_text: str, 
    api_key: str,
    max_retries: int = 3,
    retry_delay: float = 2.5
) -> Dict[str, Any]:
    """AI generation dengan fallback models ASLI Anda dan improved error handling"""
    
    if not api_key:
        return {"success": False, "error": "API Key tidak tersedia"}
    
    try:
        genai.configure(api_key=api_key)
    except Exception as e:
        return {"success": False, "error": f"Gagal konfigurasi API: {str(e)}"}
    
    generation_config = {
        "temperature": config.TEMPERATURE,
        "top_p": config.TOP_P,
        "top_k": config.TOP_K,
        "max_output_tokens": config.MAX_OUTPUT_TOKENS
    }
    
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
    ]
    
    last_error = None
    
    # MENGGUNAKAN FALLBACK MODELS ASLI ANDA
    for i, model_name in enumerate(config.FALLBACK_MODELS):
        for retry in range(max_retries):
            try:
                model = genai.GenerativeModel(model_name)
                
                with st.spinner(f"🔄 Mencoba {model_name} (percobaan {retry + 1})..."):
                    response = model.generate_content(
                        prompt_text, 
                        generation_config=generation_config, 
                        safety_settings=safety_settings
                    )
                
                # Check response validity
                if not response.candidates:
                    raise Exception("Tidak ada kandidat respons")
                
                if response.candidates[0].finish_reason.name == 'MAX_TOKENS':
                    if retry < max_retries - 1:
                        generation_config["max_output_tokens"] = int(generation_config["max_output_tokens"] * 0.8)
                        st.warning(f"Output terlalu panjang, mencoba dengan {generation_config['max_output_tokens']} tokens...")
                        continue
                    raise Exception("MAX_TOKENS_REACHED")
                
                if response.text and len(response.text.strip()) > 50:
                    return {"success": True, "content": response.text}
                else:
                    raise Exception("Respons terlalu pendek atau kosong")
                
            except Exception as e:
                last_error = str(e)
                err_str = last_error.lower()
                
                # Handle specific errors sesuai logic asli Anda
                if '429' in err_str or 'quota' in err_str or 'exhausted' in err_str:
                    if retry < max_retries - 1:
                        wait_time = retry_delay * (retry + 1)
                        st.warning(f"Rate limit. Menunggu {wait_time} detik...")
                        time.sleep(wait_time)
                        continue
                elif '404' in err_str or 'not found' in err_str:
                    # Model not found, break ke model berikutnya
                    break
                elif 'invalid' in err_str or 'authentication' in err_str:
                    return {"success": False, "error": "API Key tidak valid"}
                
                # Jika ini percobaan terakhir dari model terakhir
                if i == len(config.FALLBACK_MODELS) - 1 and retry == max_retries - 1:
                    continue
        
        # Pindah ke model berikutnya jika model saat ini gagal
        if i < len(config.FALLBACK_MODELS) - 1:
            time.sleep(1)
    
    return {"success": False, "error": f"Semua model gagal merespon. Error: {last_error}"}

# ==============================================================================
# SECTION 5: PROMPT TEMPLATES (MEMERTAHANKAN PROMPT ASLI ANDA)
# ==============================================================================

class PromptTemplates:
    """Centralized prompt templates - mempertahankan format asli Anda"""
    
    @staticmethod
    def get_notulen_prompt(combined_transcript: str) -> str:
        """Menggunakan prompt ASLI Anda dengan sedikit perbaikan struktur"""
        return f"""**INI ADALAH DATA RAPAT FORMAL PERUSAHAAN PELINDO. BUATKAN NOTULEN RAPAT DENGAN BAHASA INDONESIA YANG SANGAT FORMAL, BAKU, DAN PROFESIONAL. HANYA FOKUS PADA AGENDA, DISKUSI, DAN KEPUTUSAN SAJA.**

Anda akan menerima data transkrip/catatan rapat yang merupakan GABUNGAN dari beberapa sumber (contoh: gabungan file transkrip otomatis VTT Zoom dan catatan teks manual). Analisis dan sintesis seluruh teks gabungan tersebut sebagai satu kesatuan alur rapat yang utuh.

Buatkan notulen rapat yang rapi dan komprehensif dari data rapat berikut:

{combined_transcript}

FORMAT YANG DIHARAPKAN (Gunakan format Tabel Markdown persis seperti ini):

# Notulen Rapat

| Judul | Keterangan |
|---|---|
| Nama Rapat | [Ekstrak/Isi nama rapat] |
| Hari/Tanggal | [Ekstrak hari, tanggal] |
| Waktu | [Ekstrak waktu rapat] |
| Tempat | [Ekstrak lokasi/metode rapat] |
| Pemimpin Rapat | [Ekstrak jabatan pemimpin rapat] |
| Dibuat oleh | Group Monitoring Evaluasi Strategi Perusahaan dan Inovasi |

**Agenda:**
- [Daftar agenda rapat secara lengkap]

**Peserta Rapat:**
| No | Nama/Jabatan |
|---|---|
| 1 | [Nama peserta 1 / Jabatan] |

**Poin Diskusi dan Arahan:**
| Pembahasan / Topik | Penanggung Jawab |
|---|---|
| **[Topik Pembahasan 1 di sini]** | |
| **Poin Diskusi:** | |
| [Jabatan/Nama] menyampaikan:<br>• [Poin penyampaian 1 yang deskriptif, menjaga konteks teknis/strategis, dan menggunakan bahasa korporat formal]<br>• [Poin penyampaian 2 yang utuh dan tidak menghilangkan makna asli] | |
| [Jabatan/Nama lain] menyampaikan/menyoroti:<br>• [Poin elaborasi yang komprehensif] | |
| **Kesimpulan :** | |
| [Jabatan/Nama] memberikan arahan sebagai berikut:<br>• [Poin kesimpulan yang jelas, tegas, dan dapat ditindaklanjuti] | [Jabatan Penanggung Jawab] |

**Disclaimer:**
*Jika tidak ada tanggapan dalam tiga hari sejak dokumen ini didistribusikan, maka dokumen ini dianggap final.*

INSTRUKSI KHUSUS DAN KETAT:
1. Identifikasi Pembicara: Wajib mengekstrak siapa yang berbicara. Gunakan Jabatan (jika disebutkan/diketahui) atau nama peserta.
2. Kedalaman Makna (SANGAT PENTING): JANGAN meringkas poin diskusi terlalu ekstrem. Pertahankan substansi, detail teknis, metrik/angka, dan konteks strategis dari pembicaraan asli. Setiap bullet point (•) harus berupa kalimat atau frasa yang utuh.
3. Gaya Bahasa Profesional: Ubah bahasa lisan, slang, atau catatan kasar menjadi bahasa dokumen korporat tingkat tinggi. Gunakan kata kerja aktif/pasif yang formal (misal: "menyoroti pentingnya...", "mengusulkan skema...", "menjabarkan kendala...", "menginstruksikan agar...").
4. Struktur Diskusi & Kesimpulan: Patuhi hierarki tabel: Topik -> "Poin Diskusi:" -> Siapa menyampaikan apa -> "Kesimpulan :" -> Arahan dan Penanggung Jawab di kolom kanan.
5. Penggabungan Konteks: Gabungkan konteks dari berbagai sumber agar kronologi nyambung secara logis tanpa duplikasi informasi."""

    @staticmethod
    def get_repair_prompt(raw_notes: str) -> str:
        """Menggunakan prompt repair ASLI Anda"""
        return f"""Anda adalah asisten AI profesional untuk Pelindo. Tugas Anda adalah mereparasi draf rapat yang acak-acakan menjadi bahasa Indonesia yang sangat formal, baku, dan kaya akan konteks profesional.

DRAF KASAR:
\"\"\"
{raw_notes}
\"\"\"

FORMAT WAJIB YANG HARUS DIGUNAKAN (Gunakan Tabel Markdown):
# Notulen Rapat
| Judul | Keterangan |
|---|---|
| Nama Rapat | [Ekstrak/Buat nama rapat] |
| Hari/Tanggal | [Ekstrak tanggal] |
| Waktu | [Ekstrak waktu] |
| Tempat | [Ekstrak tempat] |
| Pemimpin Rapat | [Ekstrak pemimpin] |
| Dibuat oleh | Group Monitoring Evaluasi Strategi Perusahaan dan Inovasi |

**Agenda:**\n- [Daftar agenda]

**Peserta Rapat:**
| No | Nama/Jabatan |
|---|---|
| 1 | [Ekstrak nama] |

**Poin Diskusi dan Arahan:**
| Pembahasan / Topik | Penanggung Jawab |
|---|---|
| **[Topik Pembahasan]** | |
| **Poin Diskusi:** | |
| [Jabatan/Nama] menyampaikan:<br>• [Poin penyampaian yang telah dielevasi menjadi bahasa formal tanpa mengurangi makna asli] | |
| **Kesimpulan :** | |
| [Jabatan/Nama] memberikan arahan sebagai berikut:<br>• [Poin arahan yang tegas] | [Penanggung Jawab] |

**Disclaimer:**\n*Jika tidak ada tanggapan dalam tiga hari sejak dokumen ini didistribusikan, maka dokumen ini dianggap final.*

INSTRUKSI REPARASI:
1. Jangan menghilangkan substansi atau mengecilkan makna dari draf asli. Perbaiki tata bahasanya saja menjadi kalimat korporat yang elegan.
2. Jika ada kalimat yang terpotong di draf, buat agar terdengar masuk akal dan formal secara bisnis."""

    @staticmethod
    def get_chat_prompt(transcript: str, question: str) -> str:
        """Prompt untuk chat tanya jawab"""
        return f"""TRANSKRIP REFERENSI:
{transcript}

PERTANYAAN: {question}

INSTRUKSI: Jawablah hanya berdasarkan transkrip di atas dengan bahasa Indonesia formal.
Jika informasi tidak tersedia dalam transkrip, katakan "Tidak ditemukan dalam transkrip rapat".
Sebutkan siapa yang menyampaikan jika ada informasinya."""

# ==============================================================================
# SECTION 6: STREAMLIT UI (IMPROVED DENGAN MEMPERTAHANKAN DESAIN ASLI)
# ==============================================================================

def setup_page_config():
    """Setup Streamlit page configuration"""
    st.set_page_config(
        page_title="MNEV Intelligence | Notulen Generator", 
        page_icon="📝", 
        layout="wide",
        initial_sidebar_state="collapsed"
    )

def apply_custom_css():
    """Apply custom CSS styling (mempertahankan desain asli Anda)"""
    st.markdown("""
    <style>
        .stApp { background-color: #faf9f6; font-family: 'Inter', sans-serif; color: #292524; }
        
        /* Header Custom */
        .mnev-header {
            background: white; border-bottom: 1px solid #e5e7eb; padding: 1rem 2rem;
            display: flex; justify-content: space-between; align-items: center; margin-bottom: 2rem;
            box-shadow: 0 1px 2px 0 rgba(0, 0, 0, 0.05);
        }
        .mnev-logo-box { background: #292524; color: white; padding: 0.5rem 0.75rem; border-radius: 4px; font-weight: bold; font-size: 0.875rem; letter-spacing: 0.1em; }
        .mnev-title { font-family: 'Lora', serif; font-weight: bold; font-size: 1.25rem; color: #292524; margin: 0; display: flex; align-items: center; gap: 8px;}
        .mnev-subtitle { font-size: 0.7rem; color: #78716c; margin: 0; letter-spacing: 0.05em; font-weight: 500;}
        .mnev-badge { background: #f5f5f4; border: 1px solid #e7e5e4; padding: 2px 6px; border-radius: 4px; font-size: 0.6rem; color: #57534e; font-family: sans-serif;}
        
        /* Tombol Khusus MNEV (#596248) */
        div.stButton > button:first-child {
            background-color: #596248; color: white; width: 100%; border-radius: 8px; border: none; font-weight: 500; padding: 0.75rem; transition: all 0.2s;
        }
        div.stButton > button:first-child:hover { background-color: #4a523b; transform: translateY(-1px); }
        
        /* Container Box */
        .custom-box { background: white; padding: 1.5rem; border-radius: 1rem; box-shadow: 0 2px 10px -3px rgba(0,0,0,0.05); border: 1px solid #e7e5e4; }
        
        /* Progress bar styling */
        .stProgress > div > div { background-color: #596248; }
        
        /* Alert styling */
        .stAlert { border-radius: 8px; border-left: 4px solid #596248; }
    </style>
    """, unsafe_allow_html=True)

def render_header():
    """Render the application header"""
    st.markdown("""
    <div class="mnev-header">
        <div style="display: flex; align-items: center; gap: 1rem;">
            <div class="mnev-logo-box">MNEV</div>
            <div>
                <h1 class="mnev-title">MNEV Intelligence <span class="mnev-badge">V3.0</span></h1>
                <p class="mnev-subtitle">Notulen Generator & Repair • Group Monitoring dan Evaluasi Strategi Perusahaan</p>
            </div>
        </div>
        <div style="font-size: 0.7rem; color: #a8a29e;">⚡ AI-Powered</div>
    </div>
    """, unsafe_allow_html=True)

def initialize_session_state():
    """Initialize all session state variables"""
    defaults = {
        "transcript_text": "",
        "ai_notulen": "",
        "ai_repaired": "",
        "chat_history": [],
        "generation_in_progress": False
    }
    
    for key, default_value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

def get_api_key() -> Optional[str]:
    """Get API key from secrets or environment"""
    return st.secrets.get("api_key") or os.getenv("GOOGLE_API_KEY")

def process_transcript(uploaded_file, manual_input: str) -> str:
    """Process and combine transcript from multiple sources"""
    combined_transcript = ""
    
    if uploaded_file:
        is_valid, msg = validate_file(uploaded_file)
        if not is_valid:
            st.error(msg)
            return ""
        
        raw_text = safe_file_read(uploaded_file)
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_ext == '.vtt':
            processed = process_vtt_text(raw_text)
            combined_transcript += "SUMBER 1 (TRANSKRIP OTOMATIS):\n" + processed + "\n\n"
        else:
            combined_transcript += "SUMBER 1 (FILE TEKS):\n" + raw_text + "\n\n"
    
    if manual_input.strip():
        combined_transcript += "SUMBER 2 (CATATAN MANUAL):\n" + manual_input.strip() + "\n\n"
    
    return combined_transcript

def generate_notulen_with_progress(combined_transcript: str, api_key: str):
    """Generate notulen with progress indicators"""
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        status_text.text("📝 Menyiapkan data...")
        progress_bar.progress(10)
        
        prompt = PromptTemplates.get_notulen_prompt(combined_transcript)
        
        status_text.text("🤖 Menganalisis dan menyatukan data...")
        progress_bar.progress(30)
        
        result = generate_with_fallback(prompt, api_key)
        
        if result['success']:
            status_text.text("✅ Memproses hasil...")
            progress_bar.progress(80)
            
            st.session_state.ai_notulen = result['content']
            
            status_text.text("✨ Selesai!")
            progress_bar.progress(100)
            time.sleep(0.5)
            return True
        else:
            st.error(result['error'])
            return False
        
    finally:
        time.sleep(0.5)
        progress_bar.empty()
        status_text.empty()

# ==============================================================================
# SECTION 7: MAIN APPLICATION
# ==============================================================================

def main():
    """Main application entry point"""
    setup_page_config()
    apply_custom_css()
    render_header()
    initialize_session_state()
    
    api_key = get_api_key()
    
    if not api_key:
        st.error("🔑 API Key belum dikonfigurasi di st.secrets. Silakan tambahkan 'api_key' di .streamlit/secrets.toml")
        st.stop()
    
    # Main layout
    col_left, col_right = st.columns([4, 8], gap="large")
    
    with col_left:
        st.markdown('<div class="custom-box">', unsafe_allow_html=True)
        st.markdown("### 📄 Data Transkrip Mentah")
        st.caption("Gunakan salah satu atau gabungkan keduanya.")
        
        uploaded_file = st.file_uploader(
            "Pilih Berkas Rapat (.VTT / .TXT)", 
            type=['vtt', 'txt'],
            help="Upload file transkrip dari Zoom/Meet atau file teks biasa"
        )
        
        st.markdown("<center><span style='font-size:0.7rem; color:#a8a29e; font-weight:bold; letter-spacing:0.1em;'>DAN / ATAU</span></center>", 
                   unsafe_allow_html=True)
        
        manual_input = st.text_area(
            "Tempel transkrip manual:", 
            height=150, 
            placeholder="Teks yang diketik di sini akan otomatis digabungkan dengan file unggahan (jika ada).",
            help="Tempel catatan rapat atau teks manual di sini"
        )
        
        st.markdown("---")
        
        if st.button("🚀 Generate Notulen Otomatis", use_container_width=True):
            if not uploaded_file and not manual_input.strip():
                st.warning("⚠️ Silakan unggah file VTT atau paste teks transkrip manual.")
            else:
                combined_transcript = process_transcript(uploaded_file, manual_input)
                if combined_transcript:
                    st.session_state.transcript_text = combined_transcript
                    
                    if generate_notulen_with_progress(combined_transcript, api_key):
                        st.success("✅ Notulen berhasil digenerate!")
                        st.balloons()
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col_right:
        tab1, tab2, tab3 = st.tabs(["📝 Hasil AI Notulen", "💬 Tanya Jawab", "🛠️ Repair Notulen (Standarisasi)"])
        
        with tab1:
            if st.session_state.ai_notulen:
                # Export buttons
                c1, c2, _ = st.columns([2, 2, 6])
                with c1:
                    word_buffer = create_word_document(st.session_state.ai_notulen)
                    if word_buffer:
                        st.download_button(
                            "📄 Unduh Word (.docx)", 
                            word_buffer.getvalue(), 
                            f"Notulen_Rapat_{datetime.now().strftime('%Y%m%d')}.docx",
                            use_container_width=True
                        )
                with c2:
                    pdf_bytes = create_pdf_document(st.session_state.ai_notulen)
                    if pdf_bytes:
                        st.download_button(
                            "📕 Unduh PDF", 
                            pdf_bytes, 
                            f"Notulen_Rapat_{datetime.now().strftime('%Y%m%d')}.pdf",
                            mime="application/pdf", 
                            use_container_width=True
                        )
                    else:
                        st.error("Gagal generate PDF")
                
                st.divider()
                with st.container(border=True):
                    st.markdown(st.session_state.ai_notulen, unsafe_allow_html=True)
            else:
                st.info("💡 Belum Ada Data Transkrip. Hasil AI Notulen akan muncul di sini.")
        
        with tab2:
            if not st.session_state.transcript_text:
                st.warning("⚠️ Silakan generate notulen terlebih dahulu agar AI memiliki konteks transkrip.")
            else:
                # Display chat history with limit
                for msg in st.session_state.chat_history[-config.MAX_CHAT_HISTORY:]:
                    with st.chat_message(msg["role"]):
                        st.markdown(msg["content"])
                
                if question := st.chat_input("Tanyakan spesifik terkait transkrip..."):
                    st.session_state.chat_history.append({"role": "user", "content": question})
                    with st.chat_message("user"):
                        st.markdown(question)
                    
                    with st.chat_message("assistant"):
                        with st.spinner("💭 Menganalisis transkrip..."):
                            prompt = PromptTemplates.get_chat_prompt(
                                st.session_state.transcript_text[:15000],  # Limit context
                                question
                            )
                            response = generate_with_fallback(prompt, api_key)
                            if response['success']:
                                st.markdown(response['content'])
                                st.session_state.chat_history.append(
                                    {"role": "assistant", "content": response['content']}
                                )
                            else:
                                st.error(response['error'])
        
        with tab3:
            st.markdown("### 🔧 Standarisasi Draf Kasar")
            raw_notes = st.text_area(
                "Tempel catatan rapat yang berantakan, tulisan dari WhatsApp, atau draf kasar di sini...", 
                height=150,
                placeholder="Contoh:\nRapat hari ini membahas...\nPak Ahmad bilang...\nKesimpulan..."
            )
            
            if st.button("✨ Reparasi & Format ke Tabel", key="btn_repair", use_container_width=True):
                if not raw_notes.strip():
                    st.warning("⚠️ Silakan tempel draf kasar terlebih dahulu.")
                else:
                    prompt = PromptTemplates.get_repair_prompt(raw_notes)
                    with st.spinner("🤖 Memproses reparasi notulen..."):
                        res_rep = generate_with_fallback(prompt, api_key)
                        if res_rep['success']:
                            st.session_state.ai_repaired = res_rep['content']
                            st.success("✅ Reparasi selesai!")
                        else:
                            st.error(res_rep['error'])
            
            if st.session_state.ai_repaired:
                st.divider()
                
                # Export buttons for repaired document
                rc1, rc2, _ = st.columns([2, 2, 6])
                with rc1:
                    r_word_buffer = create_word_document(st.session_state.ai_repaired)
                    if r_word_buffer:
                        st.download_button(
                            "📄 Unduh Word", 
                            r_word_buffer.getvalue(), 
                            f"Reparasi_Notulen_{datetime.now().strftime('%Y%m%d')}.docx",
                            key="dl_rep_word", 
                            use_container_width=True
                        )
                with rc2:
                    r_pdf_bytes = create_pdf_document(st.session_state.ai_repaired)
                    if r_pdf_bytes:
                        st.download_button(
                            "📕 Unduh PDF", 
                            r_pdf_bytes, 
                            f"Reparasi_Notulen_{datetime.now().strftime('%Y%m%d')}.pdf",
                            mime="application/pdf", 
                            key="dl_rep_pdf", 
                            use_container_width=True
                        )
                
                with st.container(border=True):
                    st.markdown(st.session_state.ai_repaired, unsafe_allow_html=True)

# ==============================================================================
# RUN APPLICATION
# ==============================================================================

if __name__ == "__main__":
    main()
    

